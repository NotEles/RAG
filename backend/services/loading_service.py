from pathlib import Path
import logging
import os
from datetime import datetime
import json

logger = logging.getLogger(__name__)
"""
PDF文档加载服务类
    这个服务类提供了多种PDF文档加载方法，支持不同的加载策略和分块选项。
    主要功能：
    1. 支持多种PDF解析库：
        - PyMuPDF (fitz): 适合快速处理大量PDF文件，性能最佳
        - PyPDF: 适合简单的PDF文本提取，依赖较少
        - pdfplumber: 适合需要处理表格或需要文本位置信息的场景
        - unstructured: 适合需要更好的文档结构识别和灵活分块策略的场景
    
    2. 文档加载特性：
        - 保持页码信息
        - 支持文本分块
        - 提供元数据存储
        - 支持不同的加载策略（使用unstructured时）
 """
class LoadingService:
    """
    PDF文档加载服务类，提供多种PDF文档加载和处理方法。
    
    属性:
        total_pages (int): 当前加载PDF文档的总页数
        current_page_map (list): 存储当前文档的页面映射信息，每个元素包含页面文本和页码
    """
    
    def __init__(self):
        self.total_pages = 0
        self.current_page_map = []

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".json", ".jsonl", ".md", ".markdown", ".txt", ".csv",
        ".docx", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"
    }

    def load_document(self, file_path: str, method: str = "auto", strategy: str = None, chunking_strategy: str = None, chunking_options: dict = None) -> str:
        """Load text from supported document formats and keep a page-like map."""
        suffix = Path(file_path).suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise ValueError(f"Unsupported file type: {suffix or 'unknown'}. Supported types: {supported}")

        if suffix == ".pdf":
            pdf_method = method if method and method != "auto" else "pymupdf"
            return self.load_pdf(
                file_path,
                pdf_method,
                strategy=strategy,
                chunking_strategy=chunking_strategy,
                chunking_options=chunking_options,
            )
        if suffix in {".json", ".jsonl"}:
            return self._load_json(file_path)
        if suffix in {".md", ".markdown"}:
            return self._load_markdown(file_path)
        if suffix == ".txt":
            return self._load_text_file(file_path, source_type="text")
        if suffix == ".csv":
            return self._load_csv(file_path)
        if suffix == ".docx":
            return self._load_docx(file_path)
        if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}:
            return self._load_image(file_path)

        raise ValueError(f"Unsupported file type: {suffix}")
    
    def load_pdf(self, file_path: str, method: str, strategy: str = None, chunking_strategy: str = None, chunking_options: dict = None) -> str:
        """
        加载PDF文档的主方法，支持多种加载策略。

        参数:
            file_path (str): PDF文件路径
            method (str): 加载方法，支持 'pymupdf', 'pypdf', 'pdfplumber', 'unstructured'
            strategy (str, optional): 使用unstructured方法时的策略，可选 'fast', 'hi_res', 'ocr_only'
            chunking_strategy (str, optional): 文本分块策略，可选 'basic', 'by_title'
            chunking_options (dict, optional): 分块选项配置

        返回:
            str: 提取的文本内容
        """
        try:
            chunking_options = chunking_options or {}
            if method == "pymupdf":
                return self._load_with_pymupdf(file_path)
            elif method == "pypdf":
                return self._load_with_pypdf(file_path)
            elif method == "pdfplumber":
                return self._load_with_pdfplumber(file_path)
            elif method == "unstructured":
                return self._load_with_unstructured(
                    file_path, 
                    strategy=strategy,
                    chunking_strategy=chunking_strategy,
                    chunking_options=chunking_options
                )
            else:
                raise ValueError(f"Unsupported loading method: {method}")
        except Exception as e:
            logger.error(f"Error loading PDF with {method}: {str(e)}")
            raise

    def _set_page_map(self, blocks: list) -> str:
        self.current_page_map = [block for block in blocks if str(block.get("text", "")).strip()]
        self.total_pages = len(self.current_page_map)
        return "\n".join(block["text"] for block in self.current_page_map)

    def _load_text_file(self, file_path: str, source_type: str = "text") -> str:
        with open(file_path, "r", encoding="utf-8-sig") as f:
            text = f.read().strip()
        blocks = [{"text": text, "page": 1, "metadata": {"source_type": source_type}}] if text else []
        return self._set_page_map(blocks)

    def _load_markdown(self, file_path: str) -> str:
        text = self._read_text_with_fallback(file_path)
        sections = []
        current_title = "Markdown"
        current_lines = []

        for line in text.splitlines():
            if line.lstrip().startswith("#") and current_lines:
                sections.append({
                    "text": "\n".join(current_lines).strip(),
                    "page": len(sections) + 1,
                    "metadata": {"source_type": "markdown", "title": current_title}
                })
                current_title = line.lstrip("#").strip() or "Untitled"
                current_lines = [line]
            else:
                if line.lstrip().startswith("#"):
                    current_title = line.lstrip("#").strip() or "Untitled"
                current_lines.append(line)

        if current_lines:
            sections.append({
                "text": "\n".join(current_lines).strip(),
                "page": len(sections) + 1,
                "metadata": {"source_type": "markdown", "title": current_title}
            })

        return self._set_page_map(sections)

    def _load_json(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        if suffix == ".jsonl":
            items = []
            with open(file_path, "r", encoding="utf-8-sig") as f:
                for line in f:
                    line = line.strip()
                    if line:
                        items.append(json.loads(line))
            data = items
        else:
            with open(file_path, "r", encoding="utf-8-sig") as f:
                data = json.load(f)

        items = self._extract_json_items(data)
        blocks = []
        for idx, item in enumerate(items, 1):
            text = self._json_item_to_text(item)
            if text.strip():
                blocks.append({
                    "text": text,
                    "page": idx,
                    "metadata": {"source_type": "json", "json_index": idx}
                })
        return self._set_page_map(blocks)

    def _load_csv(self, file_path: str) -> str:
        import pandas as pd

        df = pd.read_csv(file_path)
        blocks = []
        for idx, row in df.iterrows():
            values = []
            for key, value in row.items():
                if pd.notna(value):
                    values.append(f"{key}: {value}")
            text = "\n".join(values).strip()
            if text:
                blocks.append({
                    "text": text,
                    "page": idx + 1,
                    "metadata": {"source_type": "csv", "row": idx + 1}
                })
        return self._set_page_map(blocks)

    def _load_docx(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("DOCX parsing requires python-docx. Please install it with: pip install python-docx") from exc

        document = Document(file_path)
        blocks = []
        current_lines = []
        section_title = "Document"

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = getattr(para.style, "name", "")
            if style_name.startswith("Heading") and current_lines:
                blocks.append({
                    "text": "\n".join(current_lines).strip(),
                    "page": len(blocks) + 1,
                    "metadata": {"source_type": "docx", "title": section_title}
                })
                section_title = text
                current_lines = [text]
            else:
                if style_name.startswith("Heading"):
                    section_title = text
                current_lines.append(text)

        for table_idx, table in enumerate(document.tables, 1):
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            table_text = "\n".join(rows).strip()
            if table_text:
                current_lines.append(f"Table {table_idx}:\n{table_text}")

        if current_lines:
            blocks.append({
                "text": "\n".join(current_lines).strip(),
                "page": len(blocks) + 1,
                "metadata": {"source_type": "docx", "title": section_title}
            })
        return self._set_page_map(blocks)

    def _load_image(self, file_path: str) -> str:
        try:
            from PIL import Image
            import pytesseract
        except ImportError as exc:
            raise ImportError("Image parsing requires pillow and pytesseract. Please install them with: pip install pillow pytesseract") from exc

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
        except Exception:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image).strip()

        if not text:
            raise ValueError("No text could be extracted from the image. Please check OCR/Tesseract installation or upload a clearer image.")

        return self._set_page_map([{
            "text": text,
            "page": 1,
            "metadata": {"source_type": "image", "ocr_engine": "tesseract"}
        }])

    def _read_text_with_fallback(self, file_path: str) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _extract_json_items(self, data) -> list:
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            if "chunks" in data and isinstance(data["chunks"], list):
                return data["chunks"]
            for value in data.values():
                if isinstance(value, list) and value:
                    return value
            return [data]
        return [data]

    def _json_item_to_text(self, item) -> str:
        if item is None:
            return ""
        if isinstance(item, str):
            return item.strip()
        if isinstance(item, (bool, int, float)):
            return str(item)
        if isinstance(item, list):
            parts = [self._json_item_to_text(value) for value in item]
            return "\n".join(part for part in parts if part)
        if isinstance(item, dict):
            if "content" in item and isinstance(item["content"], str):
                return item["content"].strip()
            if "question" in item and "answers" in item:
                question = str(item.get("question", "")).strip()
                answers = item.get("answers", [])
                answer_texts = []
                if isinstance(answers, list):
                    for answer in answers:
                        if isinstance(answer, dict):
                            answer_texts.append(str(answer.get("answer", "")).strip())
                        else:
                            answer_texts.append(str(answer).strip())
                answer_text = "\n".join(text for text in answer_texts if text)
                return f"问题：{question}\n回答：{answer_text}".strip()
            parts = []
            for key, value in item.items():
                if key in {"id", "chunk_id", "answer_quality"}:
                    continue
                text = self._json_item_to_text(value)
                if text:
                    parts.append(f"{key}: {text}" if not isinstance(value, (dict, list)) else text)
            return "\n".join(parts)
        return str(item).strip()
    
    def get_total_pages(self) -> int:
        """
        获取当前加载文档的总页数。

        返回:
            int: 文档总页数
        """
        return max(page_data['page'] for page_data in self.current_page_map) if self.current_page_map else 0
    
    def get_page_map(self) -> list:
        """
        获取当前文档的页面映射信息。

        返回:
            list: 包含每页文本内容和页码的列表
        """
        return self.current_page_map
    
    def _load_with_pymupdf(self, file_path: str) -> str:
        """
        使用PyMuPDF库加载PDF文档。
        适合快速处理大量PDF文件，性能最佳。

        参数:
            file_path (str): PDF文件路径

        返回:
            str: 提取的文本内容
        """
        text_blocks = []
        try:
            import fitz  # PyMuPDF

            with fitz.open(file_path) as doc:
                self.total_pages = len(doc)
                for page_num, page in enumerate(doc, 1):
                    text = page.get_text("text")
                    if text.strip():
                        text_blocks.append({
                            "text": text.strip(),
                            "page": page_num
                        })
            self.current_page_map = text_blocks
            return "\n".join(block["text"] for block in text_blocks)
        except Exception as e:
            logger.error(f"PyMuPDF error: {str(e)}")
            raise
    
    def _load_with_pypdf(self, file_path: str) -> str:
        """
        使用PyPDF库加载PDF文档。
        适合简单的PDF文本提取，依赖较少。

        参数:
            file_path (str): PDF文件路径

        返回:
            str: 提取的文本内容
        """
        try:
            from pypdf import PdfReader

            text_blocks = []
            with open(file_path, "rb") as file:
                pdf = PdfReader(file)
                self.total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_blocks.append({
                            "text": page_text.strip(),
                            "page": page_num
                        })
            self.current_page_map = text_blocks
            return "\n".join(block["text"] for block in text_blocks)
        except Exception as e:
            logger.error(f"PyPDF error: {str(e)}")
            raise
    
    def _load_with_unstructured(self, file_path: str, strategy: str = "fast", chunking_strategy: str = "basic", chunking_options: dict = None) -> str:
        """
        使用unstructured库加载PDF文档。
        适合需要更好的文档结构识别和灵活分块策略的场景。

        参数:
            file_path (str): PDF文件路径
            strategy (str): 加载策略，默认'fast'
            chunking_strategy (str): 分块策略，默认'basic'
            chunking_options (dict): 分块选项配置

        返回:
            str: 提取的文本内容
        """
        try:
            from unstructured.partition.pdf import partition_pdf

            strategy_params = {
                "fast": {"strategy": "fast"},
                "hi_res": {"strategy": "hi_res"},
                "ocr_only": {"strategy": "ocr_only"}
            }            
         
            # Prepare chunking parameters based on strategy
            chunking_params = {}
            if chunking_strategy == "basic":
                chunking_params = {
                    "max_characters": chunking_options.get("maxCharacters", 4000),
                    "new_after_n_chars": chunking_options.get("newAfterNChars", 3000),
                    "combine_text_under_n_chars": chunking_options.get("combineTextUnderNChars", 2000),
                    "overlap": chunking_options.get("overlap", 200),
                    "overlap_all": chunking_options.get("overlapAll", False)
                }
            elif chunking_strategy == "by_title":
                chunking_params = {
                    "chunking_strategy": "by_title",
                    "combine_text_under_n_chars": chunking_options.get("combineTextUnderNChars", 2000),
                    "multipage_sections": chunking_options.get("multiPageSections", False)
                }
            
            # Combine strategy parameters with chunking parameters
            params = {**strategy_params.get(strategy, {"strategy": "fast"}), **chunking_params}
            
            elements = partition_pdf(file_path, **params)
            
            # Add debug logging
            for elem in elements:
                logger.debug(f"Element type: {type(elem)}")
                logger.debug(f"Element content: {str(elem)}")
                logger.debug(f"Element dir: {dir(elem)}")
            
            text_blocks = []
            pages = set()
            
            for elem in elements:
                metadata = elem.metadata.__dict__
                page_number = metadata.get('page_number')
                
                if page_number is not None:
                    pages.add(page_number)
                    
                    # Convert element to a serializable format
                    cleaned_metadata = {}
                    for key, value in metadata.items():
                        if key == '_known_field_names':
                            continue
                        
                        try:
                            # Try JSON serialization to test if value is serializable
                            json.dumps({key: value})
                            cleaned_metadata[key] = value
                        except (TypeError, OverflowError):
                            # If not serializable, convert to string
                            cleaned_metadata[key] = str(value)
                    
                    # Add additional element information
                    cleaned_metadata['element_type'] = elem.__class__.__name__
                    cleaned_metadata['id'] = str(getattr(elem, 'id', None))
                    cleaned_metadata['category'] = str(getattr(elem, 'category', None))
                    
                    text_blocks.append({
                        "text": str(elem),
                        "page": page_number,
                        "metadata": cleaned_metadata
                    })
            
            self.total_pages = max(pages) if pages else 0
            self.current_page_map = text_blocks
            return "\n".join(block["text"] for block in text_blocks)
            
        except Exception as e:
            logger.error(f"Unstructured error: {str(e)}")
            raise
    
    def _load_with_pdfplumber(self, file_path: str) -> str:
        """
        使用pdfplumber库加载PDF文档。
        适合需要处理表格或需要文本位置信息的场景。

        参数:
            file_path (str): PDF文件路径

        返回:
            str: 提取的文本内容
        """
        text_blocks = []
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                self.total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_blocks.append({
                            "text": page_text.strip(),
                            "page": page_num
                        })
            self.current_page_map = text_blocks
            return "\n".join(block["text"] for block in text_blocks)
        except Exception as e:
            logger.error(f"pdfplumber error: {str(e)}")
            raise
    
    def save_document(self, filename: str, chunks: list, metadata: dict, loading_method: str, strategy: str = None, chunking_strategy: str = None) -> str:
        """
        保存处理后的文档数据。

        参数:
            filename (str): 原PDF文件名
            chunks (list): 文档分块列表
            metadata (dict): 文档元数据
            loading_method (str): 使用的加载方法
            strategy (str, optional): 使用的加载策略
            chunking_strategy (str, optional): 使用的分块策略

        返回:
            str: 保存的文件路径
        """
        try:
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            base_name = Path(filename).stem.split('_')[0]
            
            # Adjust the document name to include strategy if unstructured
            if loading_method == "unstructured" and strategy:
                doc_name = f"{base_name}_{loading_method}_{strategy}_{chunking_strategy}_{timestamp}"
            else:
                doc_name = f"{base_name}_{loading_method}_{timestamp}"
            
            # 构建文档数据结构，确保所有值都是可序列化的
            document_data = {
                "filename": str(filename),
                "total_chunks": int(len(chunks)),
                "total_pages": int(metadata.get("total_pages", 1)),
                "loading_method": str(loading_method),
                "loading_strategy": str(strategy) if loading_method == "unstructured" and strategy else None,
                "chunking_strategy": str(chunking_strategy) if loading_method == "unstructured" and chunking_strategy else None,
                "chunking_method": "loaded",
                "timestamp": datetime.now().isoformat(),
                "chunks": chunks
            }
            
            # 保存到文件
            filepath = os.path.join("01-loaded-docs", f"{doc_name}.json")
            os.makedirs("01-loaded-docs", exist_ok=True)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(document_data, f, ensure_ascii=False, indent=2)
                
            return filepath
            
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise
